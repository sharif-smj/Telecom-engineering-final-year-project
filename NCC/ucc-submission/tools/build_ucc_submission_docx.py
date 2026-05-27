from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Inches, Pt


ROOT = Path("/Users/sharif/telecom/final-year-project/NCC/ucc-submission")
SOURCE = ROOT / "ucc-research-support-proposal.md"
OUTPUT = ROOT / "ucc-research-support-proposal.docx"
ARCHITECTURE_IMAGE = ROOT / "assets" / "denoising-assisted-gsm-architecture.png"
MODEL_ARCHITECTURE_IMAGE = ROOT / "assets" / "dae-amc-model-architecture.png"
GANTT_TIMELINE_IMAGE = ROOT / "assets" / "project-timeline-gantt.png"

PAGE_BREAK_HEADINGS = {
    "Project Summary",
    "Introduction",
    "Methodology",
    "Budget and Budget Justification",
    "Curriculum Vitae of the Research Team",
    "References",
}

def apply_run_font(
    run,
    size: int | float = 12,
    bold: bool | None = None,
    *,
    font_name: str = "Times New Roman",
    italic: bool | None = None,
    subscript: bool | None = None,
    superscript: bool | None = None,
) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if subscript is not None:
        run.font.subscript = subscript
    if superscript is not None:
        run.font.superscript = superscript


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=110, start=110, bottom=110, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_inches: float) -> None:
    width = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths_inches: list[float]) -> None:
    widths_dxa = [int(width * 1440) for width in widths_inches]
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def add_horizontal_rule(
    doc: Document,
    *,
    before: int | float = 0,
    after: int | float = 8,
    color: str = "1F2A44",
    size: int = 8,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def clean_inline(text: str) -> str:
    text = text.replace("**", "")
    text = re.sub(r"(?<!\*)\*(?!\*)(.*?)\*(?!\*)", r"\1", text)
    return text.strip()


def section_label(text: str) -> str:
    return re.sub(r"^\d+\.\s*", "", clean_inline(text)).strip()


def should_page_break_before_heading(heading: str) -> bool:
    return section_label(heading) in PAGE_BREAK_HEADINGS


def add_runs_with_bold_label(paragraph, text: str) -> None:
    text = clean_inline(text)
    match = re.match(r"^([^:]{1,80}:)(\s*)(.*)$", text)
    if match:
        label, spacer, rest = match.groups()
        run = paragraph.add_run(label)
        apply_run_font(run, bold=True)
        if spacer:
            paragraph.add_run(spacer)
        if rest:
            add_inline_math_runs(paragraph, rest)
    else:
        add_inline_math_runs(paragraph, text)


INLINE_MATH_PATTERNS = [
    "sigma_w^2",
    "SNR_dB",
    "i_wb",
    "i_nb",
    "P_s",
    "P_w",
    "A_i",
    "f_i",
    "F_s",
    "E_theta",
    "D_theta",
    "g_phi",
    "L_DAE",
    "x_hat",
    "y_hat",
]


def add_plain_run(paragraph, text: str) -> None:
    if text:
        run = paragraph.add_run(text)
        apply_run_font(run)


def add_inline_variable(paragraph, token: str) -> None:
    def m(text: str, **kwargs) -> None:
        add_math_run(paragraph, text, **kwargs)

    if token == "sigma_w^2":
        m("σ", italic=True); m("w", subscript=True); m("2", superscript=True)
    elif token == "SNR_dB":
        m("SNR"); m("dB", subscript=True)
    elif token in {"i_wb", "i_nb"}:
        m("i", italic=True); m(token.split("_", 1)[1], subscript=True)
    elif token in {"P_s", "P_w", "A_i", "f_i", "F_s"}:
        base, sub = token.split("_", 1)
        m(base, italic=True); m(sub, subscript=True)
    elif token in {"E_theta", "D_theta", "g_phi", "L_DAE"}:
        base, sub = token.split("_", 1)
        sub = {"theta": "θ", "phi": "φ"}.get(sub, sub)
        m(base, italic=True); m(sub, subscript=True)
    elif token == "x_hat":
        m("x̂", italic=True)
    elif token == "y_hat":
        m("ŷ", italic=True)
    else:
        add_plain_run(paragraph, token)


def add_inline_math_runs(paragraph, text: str) -> None:
    idx = 0
    while idx < len(text):
        match_token = None
        for token in INLINE_MATH_PATTERNS:
            if text.startswith(token, idx):
                match_token = token
                break
        if match_token:
            add_inline_variable(paragraph, match_token)
            idx += len(match_token)
            continue
        next_positions = [
            text.find(token, idx + 1)
            for token in INLINE_MATH_PATTERNS
            if text.find(token, idx + 1) != -1
        ]
        next_idx = min(next_positions) if next_positions else len(text)
        add_plain_run(paragraph, text[idx:next_idx])
        idx = next_idx


def add_math_run(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    subscript: bool = False,
    superscript: bool = False,
    size: int | float = 12,
) -> None:
    run = paragraph.add_run(text)
    apply_run_font(
        run,
        size=size,
        bold=bold,
        font_name="Cambria Math",
        italic=italic,
        subscript=subscript,
        superscript=superscript,
    )


def add_equation(doc: Document, number: int) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    remove_table_borders(table)
    formula_cell, number_cell = table.rows[0].cells
    set_cell_width(formula_cell, 5.85)
    set_cell_width(number_cell, 0.8)
    for cell in (formula_cell, number_cell):
        set_cell_margins(cell, top=50, start=20, bottom=70, end=20)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    paragraph = formula_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = False

    def m(text: str, **kwargs) -> None:
        add_math_run(paragraph, text, **kwargs)

    if number == 1:
        m("r", italic=True); m("[n] = ")
        m("a", italic=True); m("[n]"); m("s", italic=True); m("[n] + ")
        m("i", italic=True); m("wb", subscript=True); m("[n] + ")
        m("i", italic=True); m("nb", subscript=True); m("[n] + ")
        m("w", italic=True); m("[n]")
    elif number == 2:
        m("w", italic=True); m("[n] ~ CN(0, ")
        m("σ", italic=True); m("w", subscript=True); m("2", superscript=True); m(")")
    elif number == 3:
        m("SNR"); m("dB", subscript=True); m(" = 10 log"); m("10", subscript=True)
        m("("); m("P", italic=True); m("s", subscript=True); m(" / ")
        m("P", italic=True); m("w", subscript=True); m(")")
    elif number == 4:
        m("i", italic=True); m("nb", subscript=True); m("[n] = ")
        m("A", italic=True); m("i", subscript=True); m(" cos(2π ")
        m("f", italic=True); m("i", subscript=True); m("n / ")
        m("F", italic=True); m("s", subscript=True); m(" + φ)")
    elif number == 5:
        m("a", italic=True); m("[n] = 10"); m("−L[n]/20", superscript=True)
    elif number == 6:
        m("z", italic=True); m(" = ")
        m("E", italic=True); m("θ", subscript=True); m("("); m("r", italic=True); m("), ")
        m("x̂", italic=True); m(" = ")
        m("D", italic=True); m("θ", subscript=True); m("("); m("z", italic=True); m(")")
    elif number == 7:
        m("L", italic=True); m("DAE", subscript=True); m(" = (1/N) ")
        m("∑"); m("n=1", subscript=True); m("N", superscript=True); m(" |")
        m("s", italic=True); m("[n] − "); m("x̂", italic=True); m("[n]|")
        m("2", superscript=True)
    elif number == 8:
        m("ŷ", italic=True); m(" = argmax ")
        m("g", italic=True); m("φ", subscript=True); m("("); m("x̂", italic=True); m(")")
    elif number == 9:
        m("Precision = TP / (TP + FP)")
    elif number == 10:
        m("Recall = TP / (TP + FN)")
    elif number == 11:
        m("F1 = 2(Precision × Recall) / (Precision + Recall)")
    elif number == 12:
        m("Macro F1 = average of F1 across all classes")
    else:
        m(f"Equation {number}")

    number_paragraph = number_cell.paragraphs[0]
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_paragraph.paragraph_format.space_before = Pt(3)
    number_paragraph.paragraph_format.space_after = Pt(3)
    run = number_paragraph.add_run(f"({number})")
    apply_run_font(run)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_architecture_figure(doc: Document) -> None:
    if not ARCHITECTURE_IMAGE.exists():
        raise FileNotFoundError(f"Architecture figure asset not found: {ARCHITECTURE_IMAGE}")

    figure = doc.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.space_before = Pt(8)
    figure.paragraph_format.space_after = Pt(4)
    figure.paragraph_format.keep_together = True
    figure_run = figure.add_run()
    figure_run.add_picture(str(ARCHITECTURE_IMAGE), width=Inches(6.65))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True
    run = caption.add_run("Figure 1. Denoising-assisted GSM signal processing and evaluation pipeline")
    apply_run_font(run, size=10.5, bold=True)


def add_model_architecture_figure(doc: Document) -> None:
    if not MODEL_ARCHITECTURE_IMAGE.exists():
        raise FileNotFoundError(f"Model architecture figure asset not found: {MODEL_ARCHITECTURE_IMAGE}")

    figure = doc.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.space_before = Pt(8)
    figure.paragraph_format.space_after = Pt(4)
    figure.paragraph_format.keep_together = True
    figure_run = figure.add_run()
    figure_run.add_picture(str(MODEL_ARCHITECTURE_IMAGE), width=Inches(6.65))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True
    run = caption.add_run("Figure 2. DAE-AMC denoising and classification model architecture")
    apply_run_font(run, size=10.5, bold=True)


def add_gantt_timeline_figure(doc: Document) -> None:
    if not GANTT_TIMELINE_IMAGE.exists():
        raise FileNotFoundError(f"Gantt timeline asset not found: {GANTT_TIMELINE_IMAGE}")

    figure = doc.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.space_before = Pt(8)
    figure.paragraph_format.space_after = Pt(4)
    figure.paragraph_format.keep_together = True
    figure_run = figure.add_run()
    figure_run.add_picture(str(GANTT_TIMELINE_IMAGE), width=Inches(6.65))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True
    run = caption.add_run("Figure 3. Gantt chart for the May-November 2026 implementation plan")
    apply_run_font(run, size=10.5, bold=True)


def style_paragraph(
    paragraph,
    *,
    justify: bool = True,
    space_after: int | float = 6,
    line_spacing: float = 1.0,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(space_after)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        apply_run_font(run, font_name=run.font.name or "Times New Roman")


def set_document_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(6)

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Times New Roman"
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading_1.font.size = Pt(13)
    heading_1.font.bold = True
    heading_1.paragraph_format.space_before = Pt(12)
    heading_1.paragraph_format.space_after = Pt(8)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Times New Roman"
    heading_2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading_2.font.size = Pt(12)
    heading_2.font.bold = True
    heading_2.paragraph_format.space_before = Pt(10)
    heading_2.paragraph_format.space_after = Pt(5)
    heading_2.paragraph_format.keep_with_next = True


def configure_section(section) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def add_footer_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    apply_run_font(run, size=10)


def add_centered_paragraph(doc: Document, text: str, *, size=12, bold=False, after=6, before=0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    apply_run_font(run, size=size, bold=bold)
    return paragraph


def add_main_heading(doc: Document, heading: str) -> None:
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.page_break_before = False
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(section_label(heading).upper())
    apply_run_font(run, size=13, bold=True)


def add_subheading(doc: Document, heading: str) -> None:
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(clean_inline(heading))
    apply_run_font(run, size=12, bold=True)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [clean_inline(cell.strip()) for cell in raw.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    header = rows[0]
    col_count = len(header)
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    table.autofit = False

    if col_count == 4:
        widths = [0.55, 1.0, 2.9, 2.2]
    elif col_count == 3:
        widths = [1.9, 1.15, 3.6]
    else:
        widths = [6.65 / col_count for _ in range(col_count)]
    set_table_grid(table, widths)

    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for idx, value in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = value
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "E8EEF4")
        set_cell_margins(cell, top=120, start=110, bottom=120, end=110)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            style_paragraph(paragraph, justify=False, space_after=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                apply_run_font(run, size=10, bold=True)

    for row_values in rows[1:]:
        row = table.add_row()
        set_row_cant_split(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.text = value
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell, top=110, start=110, bottom=110, end=110)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph, justify=False, space_after=0)
                if col_count == 3 and idx == 1:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif col_count == 4 and idx in {0, 1}:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    apply_run_font(run, size=10, bold=value == "Total" or row_values[0] == "Total")

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(8)


def parse_title_fields(lines: list[str], start: int) -> tuple[dict[str, str], int]:
    fields: dict[str, list[str]] = {}
    current_key: str | None = None
    i = start
    while i < len(lines):
        stripped = lines[i].strip()
        if stripped.startswith("## "):
            break
        if not stripped:
            i += 1
            continue
        match = re.match(r"^\*\*(.+?):\*\*\s*(.*)$", stripped)
        if match:
            current_key = clean_inline(match.group(1))
            value = clean_inline(match.group(2))
            fields[current_key] = [value] if value else []
        elif current_key:
            fields[current_key].append(clean_inline(stripped))
        i += 1
    return {key: "\n".join(value).strip() for key, value in fields.items()}, i


def set_title_cell_text(cell, text: str, *, bold=False, center=False, size=12) -> None:
    cell.text = ""
    for index, part in enumerate(text.split("\n")):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(part)
        apply_run_font(run, size=size, bold=bold)


def add_title_page(doc: Document, title: str, fields: dict[str, str]) -> None:
    add_centered_paragraph(doc, "RESEARCH PROPOSAL", size=15, bold=True, after=14, before=10)
    add_centered_paragraph(
        doc,
        "UGANDA COMMUNICATIONS COMMISSION (UCC)\nRESEARCH SUPPORT TO ACADEMIA",
        size=14,
        bold=True,
        after=4,
    )
    add_horizontal_rule(doc, before=2, after=12, size=10)

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    remove_table_borders(table)
    rows = [
        ("Title:", fields.get("Research study title", title)),
        ("Funding category:", fields.get("Funding category", "")),
        ("Priority research areas:", fields.get("Priority research areas", "")),
        ("Period of Performance:", fields.get("Proposed period of performance", "")),
        ("Date Prepared:", fields.get("Date prepared", "")),
        ("Principal Investigator:", fields.get("Principal Investigator", "")),
        ("Co-Principal Investigator:", fields.get("Co-Principal Investigator", "")),
        ("Student Researchers:", fields.get("Student researchers / research assistants", "")),
        ("Lead Institution:", fields.get("Lead institution", "")),
        ("Collaborating Institution:", fields.get("Collaborating institution", "")),
        ("Amount Requested:", fields.get("Total budget requested", "")),
    ]
    for label, value in rows:
        row = table.add_row()
        label_cell, value_cell = row.cells
        set_cell_width(label_cell, 2.25)
        set_cell_width(value_cell, 4.4)
        set_cell_margins(label_cell, top=66, start=30, bottom=66, end=90)
        set_cell_margins(value_cell, top=66, start=90, bottom=66, end=30)
        set_title_cell_text(label_cell, label, bold=True, size=12)
        set_title_cell_text(
            value_cell,
            value,
            bold=label in {"Title:", "Amount Requested:"},
            center=label == "Title:",
            size=12,
        )

    signature_table = doc.add_table(rows=1, cols=2)
    signature_table.autofit = False
    remove_table_borders(signature_table)
    for idx, label in enumerate(("Principal Investigator", "Co-Principal Investigator")):
        cell = signature_table.rows[0].cells[idx]
        set_cell_width(cell, 3.3)
        set_cell_margins(cell, top=105, start=80, bottom=45, end=80)
        set_title_cell_text(cell, "______________________________\n" + label, bold=False, center=True)


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = section_label(lines[0].lstrip("# ")) if lines and lines[0].startswith("# ") else ""
    doc = Document()
    configure_section(doc.sections[0])
    add_footer_page_number(doc.sections[0])
    set_document_styles(doc)

    i = 0
    previous_heading = ""
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            i += 1
            continue

        if stripped == "## 1. Title Page":
            fields, i = parse_title_fields(lines, i + 1)
            add_title_page(doc, title, fields)
            previous_heading = "1. Title Page"
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if stripped == "[ARCHITECTURE_FIGURE]":
            add_architecture_figure(doc)
            i += 1
            continue

        if stripped == "[MODEL_ARCHITECTURE_FIGURE]":
            add_model_architecture_figure(doc)
            i += 1
            continue

        if stripped == "[GANTT_TIMELINE]":
            add_gantt_timeline_figure(doc)
            i += 1
            continue

        if stripped == "[PAGE_BREAK]":
            i += 1
            continue

        equation_match = re.match(r"^Equation \((\d+)\):", stripped)
        if equation_match:
            add_equation(doc, int(equation_match.group(1)))
            i += 1
            continue

        if stripped.startswith("## "):
            heading = clean_inline(stripped[3:])
            if should_page_break_before_heading(heading):
                doc.add_page_break()
            add_main_heading(doc, heading)
            previous_heading = heading
            i += 1
            continue

        if stripped.startswith("### "):
            heading = clean_inline(stripped[4:])
            add_subheading(doc, heading)
            previous_heading = heading
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.22)
            paragraph.paragraph_format.first_line_indent = Inches(-0.22)
            add_runs_with_bold_label(paragraph, stripped)
            style_paragraph(paragraph, justify=True, space_after=5)
            i += 1
            continue

        paragraph = doc.add_paragraph()
        if "References" in previous_heading:
            paragraph.paragraph_format.left_indent = Inches(0.28)
            paragraph.paragraph_format.first_line_indent = Inches(-0.28)
            paragraph.add_run(clean_inline(stripped))
            style_paragraph(paragraph, justify=False, space_after=5)
        else:
            add_runs_with_bold_label(paragraph, stripped)
            style_paragraph(paragraph, justify=True, space_after=6)
        i += 1

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
